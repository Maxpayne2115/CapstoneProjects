from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import pandas as pd
from docx import Document
import io

app = Flask(__name__)
CORS(app)

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    project_id = data['projectId']
    project_name = data['projectName']
    client_name = data['clientName']
    applications = data['applications']
    requirements = data['requirements']
    file_type = data['fileType']

    if file_type == 'excel':
        return generate_excel(project_id, project_name, client_name, applications, requirements)
    elif file_type == 'word':
        return generate_word(project_id, project_name, client_name, applications, requirements)
    else:
        return jsonify({"error": "Invalid file type"}), 400

def generate_excel(project_id, project_name, client_name, applications, requirements):
    output = io.BytesIO()
    df = pd.DataFrame({
        "Project ID": [project_id],
        "Project Name": [project_name],
        "Client Name": [client_name],
        "Applications": [applications],
        "Requirements": [', '.join(requirements)]
    })
    df.to_excel(output, index=False)
    output.seek(0)
    return send_file(output, attachment_filename='project.xlsx', as_attachment=True)

def generate_word(project_id, project_name, client_name, applications, requirements):
    output = io.BytesIO()
    doc = Document()
    doc.add_heading('Project Details', 0)
    doc.add_paragraph(f"Project ID: {project_id}")
    doc.add_paragraph(f"Project Name: {project_name}")
    doc.add_paragraph(f"Client Name: {client_name}")
    doc.add_paragraph(f"Applications: {applications}")
    doc.add_heading('Requirements', level=1)
    for req in requirements:
        doc.add_paragraph(req, style='ListBullet')
    doc.save(output)
    output.seek(0)
    return send_file(output, attachment_filename='project.docx', as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
